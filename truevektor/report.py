"""
TrueVektor report generator.

Generates a DOCX inspection report from inspection.yaml + settings.yaml.

Usage:
    from truevektor.report import generate_report
    out_path = generate_report(folder, settings)
"""

from __future__ import annotations

import io
import warnings
from pathlib import Path
from typing import Optional

from docx.shared import RGBColor


def _fmt_date(raw: str) -> str:
    """Convert YYYY-MM-DD to Month DD, YYYY. Returns raw string unchanged on failure."""
    if not raw:
        return raw
    try:
        from datetime import datetime
        dt = datetime.strptime(raw.strip(), "%Y-%m-%d")
        return dt.strftime("%B %-d, %Y")
    except Exception:
        return raw


def _hex_to_rgb(hex_color: str) -> tuple:
    """Convert '#rrggbb' to (r, g, b) integers."""
    h = hex_color.lstrip("#")
    if len(h) != 6:
        return (255, 100, 0)
    try:
        return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except ValueError:
        return (255, 100, 0)


# ---------------------------------------------------------------------------
# Image preparation — resize & normalize into <folder>/processed/
# ---------------------------------------------------------------------------

# Target dimensions (pixels) for each image role at 200 DPI
_TARGETS = {
    "aerial":   (1300, 1300),   # 6.5" × 6.5" max
    "ortho":    (1300, 1500),   # 6.5" × 7.5" max
    "anomaly":  (600,  500),    # 2.85" × 2.4" max at ~210 DPI
}


def _process_one(src: Path, dest: Path, max_w: int, max_h: int) -> Optional[Path]:
    """
    Resize *src* to fit within max_w × max_h (no upscaling) and save as
    JPEG quality-85 at *dest*.  Skips reprocessing when *dest* is already
    newer than *src* (mtime cache).  Returns *dest* on success, None on error.
    """
    from PIL import Image

    try:
        # Cache: skip if dest already up-to-date
        if dest.exists() and dest.stat().st_mtime >= src.stat().st_mtime:
            return dest

        img = Image.open(src)
        # Flatten to RGB (handles RGBA, palette, TIFF, etc.)
        if img.mode != "RGB":
            bg = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "RGBA":
                bg.paste(img, mask=img.split()[3])
            else:
                bg.paste(img.convert("RGB"))
            img = bg

        # Only shrink, never upscale
        iw, ih = img.size
        scale = min(1.0, max_w / iw, max_h / ih)
        if scale < 1.0:
            img = img.resize((int(iw * scale), int(ih * scale)), Image.LANCZOS)

        img.save(str(dest), format="JPEG", quality=85, optimize=True)
        return dest

    except Exception as exc:
        warnings.warn(f"TrueVektor: image processing failed for {src}: {exc}")
        return None


def _prepare_images(folder: Path, data: dict) -> dict:
    """
    Process all inspection images into <folder>/processed/.

    Returns::

        {
            "aerial":    Path | None,
            "ortho":     Path | None,
            "anomalies": {str(anomaly_id): Path | None, ...},
        }
    """
    processed_dir = folder / "processed"
    processed_dir.mkdir(exist_ok=True)

    images_meta = data.get("images") or {}
    anomalies   = data.get("anomalies") or []

    result: dict = {"aerial": None, "ortho": None, "anomalies": {}}

    # Aerial / intro image
    aerial_src = _resolve_image(folder, images_meta.get("aerial_visual"))
    if aerial_src:
        result["aerial"] = _process_one(
            aerial_src,
            processed_dir / "aerial.jpg",
            *_TARGETS["aerial"],
        )

    # Orthomosaic (potentially very large)
    ortho_src = _resolve_image(folder, images_meta.get("orthomosaic"))
    if ortho_src:
        result["ortho"] = _process_one(
            ortho_src,
            processed_dir / "ortho.jpg",
            *_TARGETS["ortho"],
        )

    # Anomaly thermal images
    for anomaly in anomalies:
        aid = str(anomaly.get("id") or "")
        if not aid:
            continue
        img_val = anomaly.get("image")
        if not img_val:
            result["anomalies"][aid] = None
            continue
        src = folder / img_val  # works for both relative and absolute paths
        if not src.is_file():
            result["anomalies"][aid] = None
            continue
        result["anomalies"][aid] = _process_one(
            src,
            processed_dir / f"anomaly_{aid}.jpg",
            *_TARGETS["anomaly"],
        )

    return result


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_report(folder: Path, settings: dict) -> Path:
    """
    Build a DOCX report for the inspection at `folder`.
    Returns the path to the written .docx file inside `folder`.
    """
    import yaml
    from docx import Document
    from docx.shared import Inches, Pt

    # ── Load YAML ────────────────────────────────────────────────────────────
    with (folder / "inspection.yaml").open() as f:
        data = yaml.safe_load(f) or {}

    images_meta    = data.get("images") or {}
    inspection     = data.get("inspection") or {}
    client         = data.get("client") or {}
    property_d     = data.get("property") or {}
    tech           = data.get("technician") or {}
    equip_list     = data.get("equipment") or []
    anomalies      = data.get("anomalies") or []

    company_name   = settings.get("company_name") or ""
    logo_path_str  = settings.get("logo_path") or ""
    logo_path      = Path(logo_path_str) if logo_path_str else None
    page2_sections = settings.get("page2_sections") or []
    poly_color     = _hex_to_rgb(settings.get("polygon_stroke_color") or "#ff6400")
    poly_width     = int(settings.get("polygon_stroke_width") or 2)
    marker_color   = _hex_to_rgb(settings.get("ortho_marker_color") or "#ff6400")

    report_date    = _fmt_date(inspection.get("report_date") or inspection.get("date") or "")

    # {total_area} substitution — sum ALL anomalies regardless of polygon status
    total_area = sum(float(a.get("area_sqft") or 0) for a in anomalies)
    narrative  = (inspection.get("narrative") or "").replace(
        "{total_area}", f"{total_area:,.0f}"
    )

    # ── Pre-process images ────────────────────────────────────────────────────
    processed = _prepare_images(folder, data)

    # Build report-time anomaly list with processed image paths substituted in
    report_anomalies = []
    for a in anomalies:
        a_copy = dict(a)
        aid = str(a.get("id") or "")
        proc_path = processed["anomalies"].get(aid)
        if proc_path:
            a_copy["image"] = str(proc_path)
        report_anomalies.append(a_copy)

    # ── Create document ───────────────────────────────────────────────────────
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    try:
        sec.header_distance = Inches(0.4)
        sec.footer_distance = Inches(0.4)
    except AttributeError:
        pass

    _set_header(sec, logo_path)
    _set_footer(sec, report_date)

    # ── Page 1: Cover ─────────────────────────────────────────────────────────
    _add_cover_page(
        doc,
        facility_name=property_d.get("name") or client.get("name") or "",
        facility_address=property_d.get("address") or "",
        aerial_path=processed["aerial"],
    )

    # ── Page 2: Custom sections from settings ─────────────────────────────────
    if page2_sections:
        doc.add_page_break()
        _add_top_spacer(doc)
        for sec_def in page2_sections:
            heading = (sec_def.get("heading") or "").strip()
            body    = (sec_def.get("body") or "").strip()
            if heading:
                _add_section_heading(doc, heading)
            if body:
                for line in body.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

    # ── Page 3: Technician, Equipment, Conditions, Findings ───────────────────
    doc.add_page_break()
    _add_top_spacer(doc)
    _add_info_page(doc, inspection, tech, equip_list, narrative)

    # ── Page 4: Orthomosaic ───────────────────────────────────────────────────
    ortho_path = processed["ortho"]
    if ortho_path:
        doc.add_page_break()
        if report_anomalies:
            ortho_img = _composite_ortho(ortho_path, report_anomalies, marker_color)
            buf = _pil_to_bytes(ortho_img)
            _fit_image_from_bytes(doc, buf, ortho_path)
        else:
            _fit_image(doc, ortho_path)

    # ── Page 5+: Anomaly detail pages (2 × 2 grid) ───────────────────────────
    if report_anomalies:
        _add_anomaly_pages(doc, folder, report_anomalies, poly_color, poly_width)

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = folder / "inspection-report.docx"
    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Header / footer
# ---------------------------------------------------------------------------

HEADER_TITLE = "Thermal Inspection Report"
GRAY = RGBColor(0x88, 0x88, 0x88)
_RIGHT_TAB_POS = "9360"   # 6.5" × 1440 twips/inch


def _set_header(section, logo_path: Optional[Path]) -> None:
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    hp = header.add_paragraph()
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(0)

    pPr = hp._p.get_or_add_pPr()

    # Schema order: pBdr must come before tabs
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

    tabs_el = OxmlElement("w:tabs")
    tab_el  = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:pos"), _RIGHT_TAB_POS)
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    logo_run = hp.add_run()
    if logo_path and logo_path.is_file() and logo_path.suffix.lower() in {
        ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif"
    }:
        try:
            logo_run.add_picture(str(logo_path), height=Inches(0.25))
        except Exception:
            pass

    hp.add_run("\t")

    title_run = hp.add_run(HEADER_TITLE)
    title_run.font.size      = Pt(10)
    title_run.font.color.rgb = GRAY
    title_run.font.bold      = False


def _set_footer(section, report_date: str) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    footer = section.footer
    fp     = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after  = Pt(0)

    pPr    = fp._p.get_or_add_pPr()

    # Schema order: pBdr must come before tabs
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "CCCCCC")
    pBdr.append(top)
    pPr.append(pBdr)

    tabs_el = OxmlElement("w:tabs")
    tab_el  = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:pos"), _RIGHT_TAB_POS)
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    def _gray_run(text):
        r = fp.add_run(text)
        r.font.size      = Pt(9)
        r.font.color.rgb = GRAY
        return r

    def _gray_tab():
        """Emit an explicit <w:tab/> run — more reliable than a \t in <w:t>."""
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "888888")
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), "18")
        rpr.append(color_el)
        rpr.append(sz_el)
        r.append(rpr)
        r.append(OxmlElement("w:tab"))
        fp._p.append(r)

    date_text = f"Report Date: {report_date}" if report_date else "Report Date: —"
    _gray_run(date_text)
    _gray_tab()
    _gray_run("Page ")
    _add_field_run(fp._p, "PAGE")
    _gray_run(" of ")
    _add_field_run(fp._p, "NUMPAGES")


def _add_field_run(p_el, field_type: str) -> None:
    """Append a PAGE or NUMPAGES field code sequence to a paragraph element.

    The full begin / instrText / separate / cached-value / end sequence is
    required for Google Docs (and other non-Word renderers) to display the
    field value.  Without the 'separate' + cached run, Google Docs shows
    nothing for the field.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _run(*children):
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "888888")
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), "18")   # 9 pt = 18 half-points
        rpr.append(color_el)
        rpr.append(sz_el)
        r.append(rpr)
        for c in children:
            r.append(c)
        return r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    p_el.append(_run(begin))

    instr = OxmlElement("w:instrText")
    instr.text = f" {field_type} "
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    p_el.append(_run(instr))

    # 'separate' run + cached display value — required for Google Docs rendering
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    p_el.append(_run(sep))

    cached = OxmlElement("w:t")
    cached.text = "1"
    p_el.append(_run(cached))

    end_fc = OxmlElement("w:fldChar")
    end_fc.set(qn("w:fldCharType"), "end")
    p_el.append(_run(end_fc))


# ---------------------------------------------------------------------------
# Page 1 — Cover
# ---------------------------------------------------------------------------

def _add_cover_page(
    doc,
    facility_name: str,
    facility_address: str,
    aerial_path: Optional[Path],
) -> None:
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(72)
    spacer.paragraph_format.space_after  = Pt(0)

    heading_p = doc.add_paragraph()
    heading_p.paragraph_format.space_before = Pt(0)
    heading_p.paragraph_format.space_after  = Pt(8)
    h_run = heading_p.add_run("INSPECTION REPORT PREPARED FOR")
    h_run.font.size      = Pt(11)
    h_run.font.bold      = True
    h_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    h_run.font.all_caps  = True

    if facility_name:
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_before = Pt(0)
        name_p.paragraph_format.space_after  = Pt(2)
        n_run = name_p.add_run(facility_name)
        n_run.font.size = Pt(11)

    if facility_address:
        addr_p = doc.add_paragraph()
        addr_p.paragraph_format.space_before = Pt(0)
        addr_p.paragraph_format.space_after  = Pt(18)
        a_run = addr_p.add_run(facility_address)
        a_run.font.size = Pt(11)

    if aerial_path and aerial_path.is_file():
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(6)
        img_p.paragraph_format.space_after  = Pt(0)
        try:
            img_p.add_run().add_picture(str(aerial_path), width=Inches(6.5))
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Page 3 — Info page
# ---------------------------------------------------------------------------

def _add_info_page(
    doc,
    inspection: dict,
    tech: dict,
    equip_list: list,
    narrative: str,
) -> None:
    from docx.shared import Pt

    tech_values = [v for v in [
        tech.get("name") or "",
        tech.get("certification") or "",
        tech.get("email") or "",
        tech.get("phone") or "",
    ] if v]
    if tech_values:
        _add_section_heading(doc, "INSPECTION TECHNICIAN")
        for val in tech_values:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.add_run(val).font.size = Pt(10)

    if equip_list:
        _add_section_heading(doc, "EQUIPMENT USED")
        for eq in equip_list:
            for field in [eq.get("name"), eq.get("detector"), eq.get("temperature_range")]:
                if field and field.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after  = Pt(1)
                    p.add_run(field.strip()).font.size = Pt(10)

    cond_rows = [
        ("Inspection Date", _fmt_date(inspection.get("date") or "")),
        ("Time / Context",  inspection.get("time_context") or ""),
        ("Ambient Temp",
            f"{inspection['ambient_f']}°F" if inspection.get("ambient_f") else ""),
        ("Humidity",
            f"{inspection['rh_pct']}%" if inspection.get("rh_pct") else ""),
        ("Wind Speed",
            f"{inspection['wind_mph']} mph" if inspection.get("wind_mph") else ""),
    ]
    cond_rows = [(k, v) for k, v in cond_rows if v]
    if cond_rows:
        _add_section_heading(doc, "ENVIRONMENTAL CONDITIONS")
        _add_kv_table(doc, cond_rows, label_w_in=1.1, value_w_in=2.9)

    if narrative:
        _add_section_heading(doc, "SUMMARY OF FINDINGS")
        for para_text in narrative.split("\n"):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())


# ---------------------------------------------------------------------------
# Orthomosaic — aspect-ratio fitted image helpers
# ---------------------------------------------------------------------------

_MAX_ORTHO_W = 7.0   # inches
_MAX_ORTHO_H = 9.5   # inches


def _ortho_kwargs(img_path: Path) -> dict:
    """Return {'width': Inches(...)} or {'height': Inches(...)} to best fill page."""
    from docx.shared import Inches

    try:
        from PIL import Image as _PILImage
        with _PILImage.open(str(img_path)) as _im:
            iw, ih = _im.size
    except Exception:
        return {"width": Inches(_MAX_ORTHO_W)}

    ratio = iw / ih
    if ratio * _MAX_ORTHO_H <= _MAX_ORTHO_W:
        return {"height": Inches(_MAX_ORTHO_H)}
    return {"width": Inches(_MAX_ORTHO_W)}


def _fit_image(doc, img_path: Path) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    try:
        p.add_run().add_picture(str(img_path), **_ortho_kwargs(img_path))
    except Exception:
        pass


def _fit_image_from_bytes(doc, buf, ref_path: Path) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    try:
        p.add_run().add_picture(buf, **_ortho_kwargs(ref_path))
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Anomaly grid pages
# ---------------------------------------------------------------------------

def _add_anomaly_pages(
    doc,
    folder: Path,
    anomalies: list,
    poly_color=(255, 100, 0),
    poly_width=2,
) -> None:
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    COLS       = 2
    IMG_W      = Inches(2.85)
    CELL_W     = Inches(3.25)   # 4680 DXA — two cells fill full 9360 DXA content width
    CELL_W_DXA = 4680

    pages = [anomalies[i:i + 4] for i in range(0, len(anomalies), 4)]

    for page_idx, page_anomalies in enumerate(pages):
        padded = page_anomalies + [None] * (4 - len(page_anomalies))

        # Start a new page and leave 1" of blank space at the top of the body.
        # Using page_break_before on the spacer paragraph avoids a stray empty
        # page-break paragraph that can land on its own page when the preceding
        # content fills the page completely.
        sp = doc.add_paragraph()
        sp.paragraph_format.page_break_before = True
        sp.paragraph_format.space_before      = Inches(1.0)
        sp.paragraph_format.space_after       = 0

        tbl = doc.add_table(rows=2, cols=COLS)
        tbl.style = "Table Grid"
        _remove_table_borders(tbl)
        _set_table_width(tbl, CELL_W_DXA * COLS)   # 9360 DXA

        for idx, anomaly in enumerate(padded):
            row_i = idx // COLS
            col_i = idx  % COLS
            cell  = tbl.cell(row_i, col_i)
            cell.width = CELL_W
            # Row 0: normal top, generous bottom; Row 1: extra top for inter-row gap
            if row_i == 0:
                _set_cell_margins(cell, top_dxa=240, bottom_dxa=480)
            else:
                _set_cell_margins(cell, top_dxa=960, bottom_dxa=240)

            if anomaly is None:
                continue

            img_rel     = anomaly.get("image")
            poly_coords = anomaly.get("polygon_coords")
            anomaly_id  = anomaly.get("id") or str(idx + 1)
            location    = anomaly.get("location") or ""
            area        = anomaly.get("area_sqft") or ""
            observation = anomaly.get("observation") or ""

            img_bytes = None
            if img_rel:
                img_path = folder / img_rel   # abs path works fine here
                if img_path.is_file():
                    try:
                        composited = _composite_anomaly(img_path, poly_coords, poly_color, poly_width)
                        img_bytes  = _pil_to_bytes(composited)
                    except Exception:
                        img_bytes = None
                    if img_bytes is None:
                        try:
                            img_bytes = io.BytesIO(img_path.read_bytes())
                        except Exception:
                            img_bytes = None

            ip = cell.paragraphs[0]
            ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if img_bytes:
                try:
                    ip.add_run().add_picture(img_bytes, width=IMG_W)
                except Exception:
                    ip.add_run("[Image unavailable]")
            else:
                ip.add_run("[No image]")

            def _detail(label, value, bold_label=True):
                dp = cell.add_paragraph()
                run_l = dp.add_run(f"{label}: ")
                run_l.font.bold = bold_label
                run_l.font.size = Pt(9)
                run_v = dp.add_run(str(value))
                run_v.font.size = Pt(9)
                dp.paragraph_format.space_before = Pt(2)
                dp.paragraph_format.space_after  = Pt(2)

            _detail("Anomaly", f"#{anomaly_id}")
            if location:
                _detail("Location", location)
            if area:
                _detail("Area", f"{float(area):,.0f} sq ft")
            if observation:
                _detail("Observation", observation)


# ---------------------------------------------------------------------------
# Image compositing (Pillow)
# ---------------------------------------------------------------------------

def _composite_anomaly(img_path: Path, poly_coords, stroke_color=(255, 100, 0), stroke_width=2):
    """Draw polygon outline on a thermal image."""
    from PIL import Image, ImageDraw

    img = Image.open(img_path).convert("RGBA")
    w, h = img.size

    overlay = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw    = ImageDraw.Draw(overlay)

    if poly_coords and len(poly_coords) >= 3:
        pixels = [(float(x) * w, float(y) * h) for x, y in poly_coords]
        r, g, b = stroke_color
        line_w  = max(stroke_width, w // 300)
        for i in range(len(pixels)):
            draw.line(
                [pixels[i], pixels[(i + 1) % len(pixels)]],
                fill=(r, g, b, 230),
                width=line_w,
            )

    return Image.alpha_composite(img, overlay).convert("RGB")


def _contrast_text_color(r, g, b):
    """Return black or white (255 alpha) depending on which contrasts better with (r,g,b)."""
    luminance = 0.299 * r + 0.587 * g + 0.114 * b
    return (0, 0, 0, 255) if luminance > 140 else (255, 255, 255, 255)


def _composite_ortho(ortho_path: Path, anomalies: list, marker_color=(220, 40, 40)):
    """Draw numbered circle pins on the orthomosaic."""
    from PIL import Image, ImageDraw, ImageFont

    img = Image.open(ortho_path).convert("RGBA")
    w, h = img.size

    overlay = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw    = ImageDraw.Draw(overlay)

    # Larger markers — scale with image size, minimum 24px radius
    radius    = max(24, min(w, h) // 40)
    font_size = max(18, int(radius * 1.1))

    font = None
    for font_path in [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",
        "/Library/Fonts/Arial Bold.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]:
        try:
            font = ImageFont.truetype(font_path, font_size)
            break
        except Exception:
            continue

    r, g, b = marker_color
    text_color   = _contrast_text_color(r, g, b)
    # High-contrast outline color: opposite of text color
    outline_color = (0, 0, 0, 200) if text_color[0] == 255 else (255, 255, 255, 200)
    outline_w    = max(2, radius // 10)

    for anomaly in anomalies:
        marker = anomaly.get("orthomosaic_marker")
        if not marker:
            continue
        mx    = float(marker.get("x") or 0) * w
        my    = float(marker.get("y") or 0) * h
        label = str(anomaly.get("id") or "?")

        # Circle with contrasting outline
        draw.ellipse(
            [mx - radius, my - radius, mx + radius, my + radius],
            fill=(r, g, b, 230),
            outline=outline_color,
            width=outline_w,
        )

        if font:
            bbox = draw.textbbox((0, 0), label, font=font)
            tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        else:
            tw, th = font_size * len(label) * 0.6, font_size
        tx = mx - tw / 2
        ty = my - th / 2
        draw.text((tx, ty), label, fill=text_color, font=font)

    return Image.alpha_composite(img, overlay).convert("RGB")


def _pil_to_bytes(img) -> io.BytesIO:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _resolve_image(folder: Path, rel_or_abs: Optional[str]) -> Optional[Path]:
    """Resolve a relative or absolute image path; return Path if file exists, else None."""
    if not rel_or_abs:
        return None
    p = folder / rel_or_abs   # Python handles absolute paths correctly here
    return p if p.is_file() else None


def _add_top_spacer(doc) -> None:
    """Add half an inch of blank space at the top of a page body."""
    from docx.shared import Inches
    _add_top_spacer_in(doc, Inches(0.5))


def _add_top_spacer_in(doc, height) -> None:
    """Add a blank spacer paragraph of the given height (docx length value)."""
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = height
    sp.paragraph_format.space_after  = 0


def _add_section_heading(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

    run = p.add_run(text)
    run.font.bold      = True
    run.font.size      = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.all_caps  = True

    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_kv_table(doc, rows: list,
                  label_w_in: float = 1.8, value_w_in: float = 4.7) -> None:
    from docx.shared import Inches, Pt

    if not rows:
        return

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    _remove_table_borders(tbl)
    total_dxa = int((label_w_in + value_w_in) * 1440)
    _set_table_width(tbl, total_dxa)

    for i, (label, value) in enumerate(rows):
        lc = tbl.cell(i, 0)
        vc = tbl.cell(i, 1)
        lc.width = Inches(label_w_in)
        vc.width = Inches(value_w_in)

        run_l = lc.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.size = Pt(10)

        vc.paragraphs[0].add_run(value).font.size = Pt(10)


def _remove_table_borders(tbl) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Table-level borders
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)

    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)

    # Cell-level borders — override whatever the table style defines
    for row in tbl.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            existing = tc_pr.find(qn("w:tcBorders"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none")
                tc_borders.append(el)
            tc_pr.append(tc_borders)


def _set_cell_margins(cell, top_dxa: int = 0, bottom_dxa: int = 0,
                      left_dxa: int = 0, right_dxa: int = 0) -> None:
    """Set internal padding on a table cell in DXA units."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top_dxa), ("bottom", bottom_dxa),
                      ("left", left_dxa), ("right", right_dxa)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            tcMar.append(el)
    tcPr.append(tcMar)


def _set_table_width(tbl, width_dxa: int) -> None:
    """Set an explicit table width in DXA units.

    Google Docs (and other non-Word renderers) require an explicit w:tblW on
    the table element to know how wide to render it.  Without this, tables
    may collapse to zero width and their content appears invisible.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)

    # Remove any existing tblW
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    # Insert after tblStyle (if present) but before tblBorders — schema order
    children = list(tbl_pr)
    insert_after = -1
    for i, child in enumerate(children):
        if child.tag == qn("w:tblStyle"):
            insert_after = i
    tbl_pr.insert(insert_after + 1, tbl_w)
